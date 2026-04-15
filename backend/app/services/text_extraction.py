"""Text extraction from uploaded documents.

Supported formats: PDF (PyMuPDF), DOCX (python-docx), TXT, MD.
PPT/PPTX would require python-pptx — not installed yet.
"""
from __future__ import annotations

import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class UnsupportedFormat(Exception):
    """Raised when the file type isn't supported by the extractor."""


def extract_text(file_path: str | Path) -> str:
    """Dispatch to the right extractor based on file extension.

    Returns the full text content as a single string. Paragraph breaks are
    preserved as `\n\n` so the chunker can normalize them.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File does not exist: {path}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext in {".docx"}:
        return _extract_docx(path)
    if ext in {".txt", ".md"}:
        return _extract_text_file(path)
    raise UnsupportedFormat(f"Cannot extract text from '{ext}' files yet")


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF using PyMuPDF."""
    parts: list[str] = []
    try:
        with fitz.open(str(path)) as doc:
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text("text")
                if page_text.strip():
                    parts.append(page_text.strip())
    except Exception as e:
        logger.exception("Failed to extract PDF text: %s", path)
        raise RuntimeError(f"PDF extraction failed: {e}") from e

    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    """Extract text from a .docx using python-docx."""
    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        logger.exception("Failed to open DOCX: %s", path)
        raise RuntimeError(f"DOCX extraction failed: {e}") from e

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also pull text from tables (simple flatten)
    for table in doc.tables:
        for row in table.rows:
            row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_parts:
                parts.append(" | ".join(row_parts))

    return "\n\n".join(parts)


def _extract_text_file(path: Path) -> str:
    """Read a plain text or markdown file."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1", errors="replace")
